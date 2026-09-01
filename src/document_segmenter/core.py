from pathlib import Path

from .models import Segment


def _markdown_segments(text: str, source: str) -> list[Segment]:
    segments: list[Segment] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            segments.append(Segment(" ".join(paragraph), source=source))
            paragraph.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if line.startswith("#"):
            prefix, _, heading = line.partition(" ")
            if heading and set(prefix) == {"#"}:
                flush()
                level = len(prefix)
                segments.append(
                    Segment(heading.strip(), "title" if level == 1 else "subtitle", level, source)
                )
                continue
        paragraph.append(line)
    flush()
    return segments


def _segment_docx(path: Path) -> list[Segment]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install DOCX support with: pip install -e .[documents]") from exc

    document = Document(path)
    segments: list[Segment] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading"):
            try:
                level = int(style.rsplit(" ", 1)[-1])
            except ValueError:
                level = 2
            kind = "title" if level == 1 else "subtitle"
            segments.append(Segment(text, kind, level, path.name))
        else:
            segments.append(Segment(text, source=path.name))
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            segments.append(Segment("\n".join(rows), "table", source=path.name))
    return segments


def _segment_pdf(path: Path) -> list[Segment]:
    try:
        import pymupdf4llm
    except ImportError as exc:
        raise RuntimeError("Install PDF support with: pip install -e .[documents]") from exc
    return _markdown_segments(pymupdf4llm.to_markdown(path), path.name)


def segment_file(file_path: str | Path) -> list[Segment]:
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _markdown_segments(path.read_text(encoding="utf-8"), path.name)
    if suffix == ".docx":
        return _segment_docx(path)
    if suffix == ".pdf":
        return _segment_pdf(path)
    raise ValueError(f"Unsupported file format: {suffix or '<none>'}")
